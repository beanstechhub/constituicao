"""Extração de texto de PDF, DOCX, TXT com fallback OCR para PDFs corrompidos."""
from __future__ import annotations
import base64
import logging
import os
from pathlib import Path
from typing import IO

import requests

log = logging.getLogger(__name__)

OCR_URL = os.environ.get("OCR_SERVICE_URL", "https://ocr-service-d7id7njwla-rj.a.run.app")


def texto_legivel(texto: str) -> bool:
    if len(texto.strip()) < 30:
        return False
    alpha = sum(1 for c in texto if c.isalpha())
    return alpha / max(len(texto), 1) > 0.4


def extrair_pdf_ocr(pdf_bytes: bytes, filename: str) -> str:
    resp = requests.post(f"{OCR_URL}/ocr", json={
        "file": base64.b64encode(pdf_bytes).decode(),
        "filename": filename
    }, timeout=60)
    resp.raise_for_status()
    return resp.json()["text"]


def extrair_pdf(stream: IO[bytes], filename: str = "document.pdf") -> str:
    import pypdf
    pdf_bytes = stream.read()
    stream.seek(0)
    reader = pypdf.PdfReader(stream)
    partes = [(page.extract_text() or "") for page in reader.pages]
    texto = "\n".join(partes)

    if texto_legivel(texto):
        return texto

    log.info("pypdf extraiu texto ilegível (%d chars), tentando OCR...", len(texto))
    try:
        return extrair_pdf_ocr(pdf_bytes, filename)
    except Exception as e:
        log.warning("OCR fallback falhou: %s — retornando texto pypdf", e)
        return texto


def extrair_docx(stream: IO[bytes]) -> str:
    from docx import Document
    doc = Document(stream)
    partes: list[str] = []
    for p in doc.paragraphs:
        partes.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    partes.append(cell.text)
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text.strip():
                partes.append(p.text)
        for p in section.footer.paragraphs:
            if p.text.strip():
                partes.append(p.text)
    return "\n".join(partes)


def extrair(filename: str, stream: IO[bytes]) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extrair_pdf(stream, filename)
    if ext == ".docx":
        return extrair_docx(stream)
    if ext == ".txt":
        return stream.read().decode("utf-8", errors="replace")
    raise ValueError(f"Formato não suportado: {ext}")
